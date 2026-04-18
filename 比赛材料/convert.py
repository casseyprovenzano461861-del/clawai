import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_line_spacing(p, line_val='360'):
    pPr = p._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), line_val)
    spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)

def add_heading(doc, text, level):
    """用 add_heading 保留样式，再对 run 设置字体颜色等"""
    p = doc.add_heading(text, level=level)
    # 对已有 run 设置样式，不用 clear()
    color_map = {
        1: RGBColor(0x1F, 0x35, 0x64),
        2: RGBColor(0x1F, 0x35, 0x64),
        3: RGBColor(0x2E, 0x74, 0xB5),
        4: RGBColor(0x40, 0x40, 0x40),
    }
    size_map = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(12)}
    align_map = {1: WD_ALIGN_PARAGRAPH.CENTER, 2: WD_ALIGN_PARAGRAPH.LEFT}

    for run in p.runs:
        run.font.name = '黑体'
        run.font.size = size_map.get(level, Pt(12))
        run.font.bold = True
        run.font.color.rgb = color_map.get(level, RGBColor(0x1F, 0x35, 0x64))
        run._r.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), '黑体')

    if level in align_map:
        p.alignment = align_map[level]
    return p

def insert_toc_at_beginning(doc):
    body = doc.element.body

    # 分页符
    pgbrk_p = OxmlElement('w:p')
    pgbrk_r = OxmlElement('w:r')
    pgbrk_br = OxmlElement('w:br')
    pgbrk_br.set(qn('w:type'), 'page')
    pgbrk_r.append(pgbrk_br)
    pgbrk_p.append(pgbrk_r)

    # TOC 域
    toc_p = OxmlElement('w:p')
    toc_r = OxmlElement('w:r')
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    toc_r.append(fldChar_begin)
    toc_r.append(instrText)
    toc_r.append(fldChar_separate)
    toc_r.append(fldChar_end)
    toc_p.append(toc_r)

    # 目录标题
    title_p = OxmlElement('w:p')
    title_pPr = OxmlElement('w:pPr')
    title_jc = OxmlElement('w:jc')
    title_jc.set(qn('w:val'), 'center')
    title_pPr.append(title_jc)
    title_p.append(title_pPr)
    title_r = OxmlElement('w:r')
    title_rPr = OxmlElement('w:rPr')
    title_b = OxmlElement('w:b')
    title_sz = OxmlElement('w:sz')
    title_sz.set(qn('w:val'), '28')
    title_szCs = OxmlElement('w:szCs')
    title_szCs.set(qn('w:val'), '28')
    title_rPr.append(title_b)
    title_rPr.append(title_sz)
    title_rPr.append(title_szCs)
    title_r.append(title_rPr)
    title_t = OxmlElement('w:t')
    title_t.text = '目  录'
    title_r.append(title_t)
    title_p.append(title_r)

    body.insert(0, pgbrk_p)
    body.insert(0, toc_p)
    body.insert(0, title_p)

def md_to_docx(md_path, docx_path):
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(3.17)

    normal = doc.styles['Normal']
    normal.font.name = '宋体'
    normal.font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        if line.startswith('#### '):
            add_heading(doc, line[5:], 4)
        elif line.startswith('### '):
            add_heading(doc, line[4:], 3)
        elif line.startswith('## '):
            add_heading(doc, line[3:], 2)
        elif line.startswith('# '):
            add_heading(doc, line[2:], 1)
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            data_lines = [l for l in table_lines if not re.match(r'^\|[\s\-\|:]+\|$', l)]
            if data_lines:
                cols = [c.strip() for c in data_lines[0].split('|')[1:-1]]
                table = doc.add_table(rows=1, cols=len(cols))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for j, col in enumerate(cols):
                    cell = table.rows[0].cells[j]
                    cell.text = col
                    set_cell_bg(cell, '1F3564')
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.font.bold = True
                            run.font.name = '黑体'
                            run.font.size = Pt(11)
                for idx, row_line in enumerate(data_lines[1:]):
                    cells_data = [c.strip() for c in row_line.split('|')[1:-1]]
                    row = table.add_row()
                    bg = 'EEF2F7' if idx % 2 == 0 else 'FFFFFF'
                    for j, cell_text in enumerate(cells_data):
                        if j < len(row.cells):
                            row.cells[j].text = cell_text
                            set_cell_bg(row.cells[j], bg)
                            for para in row.cells[j].paragraphs:
                                for run in para.runs:
                                    run.font.size = Pt(11)
                doc.add_paragraph('')
            continue
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
            set_line_spacing(p)
        elif re.match(r'^\d+\. ', line):
            p = doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
            p.paragraph_format.space_after = Pt(4)
            set_line_spacing(p)
        elif line.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            if code_lines:
                p = doc.add_paragraph('\n'.join(code_lines))
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after  = Pt(4)
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
        elif line.strip() == '':
            doc.add_paragraph('')
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            set_line_spacing(p)
            for part in re.split(r'(\*\*[^*]+\*\*)', line):
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = '黑体'
                    run.font.size = Pt(12)
                else:
                    run = p.add_run(part)
                    run.font.name = '宋体'
                    run.font.size = Pt(12)
        i += 1

    insert_toc_at_beginning(doc)

    # 页脚页码
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.clear()
    run = fp.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText2 = OxmlElement('w:instrText')
    instrText2.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText2)
    run._r.append(fldChar2)
    run.font.size = Pt(10)

    doc.save(docx_path)
    print(f"已生成: {os.path.basename(docx_path)}")

base = "e:/ClawAI/比赛材料"
for f in ["项目概要介绍.md", "项目详细方案.md", "项目详细分工及过程文档.md", "安装部署文档.md", "核心功能模块说明.md"]:
    md_to_docx(os.path.join(base, f), os.path.join(base, f.replace('.md', '.docx')))
